import streamlit as st
from openai import OpenAI
import re
import pandas as pd
from io import BytesIO
from datetime import datetime, date
import plotly.graph_objects as go
import plotly.express as px
from streamlit import cache_data, cache_resource
import streamlit.components.v1 as components
import json

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import base64
from pathlib import Path

@cache_resource
def get_school_badge():
    path = Path(__file__).parent / "assets" / "logo.jpg"
    if path.exists():
        with open(path, "rb") as f:
            return "data:image/jpeg;base64," + base64.b64encode(f.read()).decode()
    return None

# ═══════════════════════════════════════════════════════════════════════
# API 配置
# ═══════════════════════════════════════════════════════════════════════
API_KEY = st.secrets["DEEPSEEK_API_KEY"]
client = OpenAI(api_key=API_KEY, base_url="https://api.deepseek.com")

st.set_page_config(
    page_title="校园财务问答助手",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ═══════════════════════════════════════════════════════════════════════
# 全局样式（现代后台管理系统主题）
# ═══════════════════════════════════════════════════════════════════════
@cache_resource
def get_css():
    return """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@300;400;500;600;700&display=swap');

    :root {
        --primary: #1F4E8C;
        --primary-light: #E8F0FA;
        --primary-dark: #153A6A;
        --accent: #2BB3B1;
        --accent-light: #E6F8F8;
        --bg-body: #F6FAFC;
        --bg-card: #FFFFFF;
        --text: #1A2C3C;
        --text-muted: #7A8FA0;
        --text-light: #A0B4C4;
        --border: #E5EEF3;
        --shadow: 0 2px 12px rgba(31,78,140,0.06);
        --shadow-hover: 0 6px 24px rgba(31,78,140,0.10);
        --radius: 12px;
        --radius-sm: 8px;
        --risk-low: #27AE60;
        --risk-med: #F39C12;
        --risk-high: #E74C3C;
    }

    /* ===== 全局基础 ===== */
    .stApp { background: var(--bg-body); color: var(--text); font-family: 'Noto Sans SC', sans-serif; }
    header[data-testid="stHeader"] { background: rgba(246,250,252,0.95); backdrop-filter: blur(6px); border-bottom: 1px solid var(--border); display: none; }
    div[data-testid="stStatusWidget"] { visibility: hidden; height: 0%; position: fixed; }
    .main > div { padding: 1.5rem 2rem 2rem; }
    section[data-testid="stSidebar"] > div { padding-top: 0; }

    /* ===== 侧边栏导航 ===== */
    section[data-testid="stSidebar"] {
        background: #FFFFFF;
        border-right: 1px solid var(--border);
        min-width: 240px !important;
        max-width: 240px !important;
        box-shadow: 2px 0 16px rgba(31,78,140,0.04);
    }
    .sidebar-logo {
        text-align: center;
        padding: 24px 16px 18px;
        border-bottom: 1px solid var(--border);
        margin-bottom: 10px;
    }
    .sidebar-badge {
        width: 52px; height: 52px; border-radius: 50%;
        background: linear-gradient(135deg, var(--primary) 0%, var(--accent) 100%);
        display: flex; align-items: center; justify-content: center;
        margin: 0 auto 10px;
        font-size: 20px; color: white; font-weight: 700;
        box-shadow: 0 3px 10px rgba(31,78,140,0.2);
    }
    .sidebar-badge img { width: 52px; height: 52px; border-radius: 50%; object-fit: cover; }
    .sidebar-sysname { font-weight: 700; font-size: 0.92rem; color: var(--primary); line-height: 1.4; }
    .sidebar-subtitle { font-size: 0.65rem; color: var(--text-light); margin-top: 2px; letter-spacing: 1px; }
    .sidebar-school { font-size: 0.62rem; color: var(--text-muted); margin-top: 4px; }

    /* 导航菜单 */
    .nav-section-label {
        font-size: 0.62rem; color: var(--text-light); text-transform: uppercase;
        letter-spacing: 1.5px; padding: 12px 16px 4px; font-weight: 600;
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] > div {
        flex-direction: column; gap: 0; padding: 0 8px;
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label {
        display: flex; align-items: center;
        padding: 10px 14px; margin: 1px 0;
        border-radius: var(--radius-sm);
        cursor: pointer; font-size: 0.85rem; font-weight: 500;
        transition: all 0.15s;
        color: var(--text);
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label:hover {
        background: var(--primary-light);
        color: var(--primary);
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label[data-selected="true"] {
        background: var(--primary-light) !important;
        color: var(--primary) !important;
        font-weight: 600;
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label[data-selected="true"]::before {
        content: "▎";
        font-size: 1.2rem; font-weight: 700;
        color: var(--primary);
        margin-right: 2px; margin-left: -6px;
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] input { display: none; }
    .sidebar-footer {
        position: fixed; bottom: 0; width: 240px;
        text-align: center; padding: 12px 16px;
        border-top: 1px solid var(--border);
        background: #FFFFFF;
    }
    .sidebar-footer-text { font-size: 0.6rem; color: var(--text-light); }

    /* ===== 卡片系统 ===== */
    .card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 20px 24px;
        margin-bottom: 18px;
        box-shadow: var(--shadow);
        transition: box-shadow 0.2s;
    }
    .card:hover { box-shadow: var(--shadow-hover); transform: translateY(-1px); }
    .card-title {
        font-size: 1rem; font-weight: 600; color: var(--text);
        margin-bottom: 14px; padding-bottom: 10px;
        border-bottom: 1px solid var(--border);
        display: flex; align-items: center; gap: 8px;
    }

    /* ===== Hero ===== */
    .hero {
        background: linear-gradient(135deg, #FFFFFF 0%, var(--primary-light) 60%, var(--accent-light) 100%);
        border: 1px solid var(--border);
        border-radius: 16px;
        padding: 36px 40px;
        margin-bottom: 24px;
        position: relative; overflow: hidden;
    }
    .hero::after {
        content: "";
        position: absolute; top: -40%; right: -10%;
        width: 260px; height: 260px;
        background: radial-gradient(circle, rgba(43,179,177,0.08) 0%, transparent 70%);
        border-radius: 50%;
    }
    .hero-content { position: relative; z-index: 1; }
    .hero-title { font-size: 2rem; font-weight: 700; color: var(--primary); letter-spacing: 2px; margin: 0 0 6px; }
    .hero-sub { color: var(--text-muted); font-size: 0.85rem; margin: 0 0 14px; font-weight: 300; letter-spacing: 0.5px; }
    .hero-badge {
        display: inline-block;
        background: linear-gradient(135deg, var(--primary) 0%, var(--accent) 100%);
        color: white; border-radius: 16px; padding: 3px 14px;
        font-size: 0.7rem; font-weight: 500;
        margin: 2px 5px 2px 0;
    }
    .hero-school-badge {
        width: 130px; height: 130px; border-radius: 50%;
        background: transparent;
        display: flex; align-items: center; justify-content: center;
        box-shadow: 0 6px 28px rgba(0,0,0,0.12), 0 2px 8px rgba(0,0,0,0.06);
        flex-shrink: 0;
        font-size: 24px; color: white; font-weight: 700;
    }
    .hero-school-badge img { width: 130px; height: 130px; border-radius: 50%; object-fit: cover; }

    /* ===== 能力指标卡 ===== */
    .cap-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 14px; margin: 16px 0; }
    .cap-card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 20px 16px;
        text-align: center;
        box-shadow: var(--shadow);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .cap-card:hover { transform: translateY(-3px); box-shadow: var(--shadow-hover); }
    .cap-icon { font-size: 1.6rem; margin-bottom: 8px; }
    .cap-value { font-size: 1.1rem; font-weight: 700; color: var(--primary); }
    .cap-label { font-size: 0.78rem; color: var(--text-muted); margin-top: 2px; }

    /* ===== 指标卡(小) ===== */
    .mini-cap-card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius-sm);
        padding: 12px 10px;
        text-align: center;
    }
    .mini-cap-icon { font-size: 1.2rem; }
    .mini-cap-label { font-size: 0.72rem; color: var(--text-muted); font-weight: 500; }

    /* ===== 痛点/方案卡片 ===== */
    .pain-card {
        background: var(--bg-card);
        border: 1px solid #F5E0E0;
        border-radius: var(--radius-sm);
        padding: 14px 16px;
        border-left: 4px solid var(--risk-high);
        box-shadow: var(--shadow);
        margin-bottom: 8px;
        transition: transform 0.2s, box-shadow 0.2s, border-left-width 0.2s;
    }
    .pain-card:hover {
        transform: translateX(4px);
        box-shadow: 0 4px 16px rgba(231,76,60,0.12);
        border-left-width: 6px;
    }
    .solve-card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius-sm);
        padding: 14px 16px;
        border-left: 4px solid var(--accent);
        box-shadow: var(--shadow);
        margin-bottom: 8px;
        transition: transform 0.2s, box-shadow 0.2s, border-left-width 0.2s;
    }
    .solve-card:hover {
        transform: translateX(4px);
        box-shadow: 0 4px 16px rgba(43,179,177,0.12);
        border-left-width: 6px;
    }

    /* ===== 对比表格 ===== */
    .compare-table { width: 100%; border-collapse: collapse; border-radius: var(--radius); overflow: hidden; font-size: 0.85rem; }
    .compare-table th { background: var(--primary); color: white; padding: 11px 16px; text-align: center; font-weight: 600; }
    .compare-table td { padding: 10px 16px; border-bottom: 1px solid var(--border); background: var(--bg-card); }
    .compare-table tr:last-child td { border-bottom: none; }
    .compare-table td:first-child { font-weight: 600; color: var(--text-muted); white-space: nowrap; }
    .compare-table td:nth-child(2) { color: #999; }
    .compare-table td:last-child { color: var(--accent); font-weight: 500; }

    /* ===== 推荐卡片 ===== */
    .rec-card {
        background: var(--primary-light);
        border: 1px solid rgba(31,78,140,0.15);
        border-radius: var(--radius);
        padding: 16px 20px;
    }
    .rec-title { font-weight: 700; color: var(--primary); font-size: 0.95rem; margin-bottom: 8px; }
    .rec-item {
        font-size: 0.84rem; color: var(--text); margin: 4px 0; padding: 4px 8px 4px 12px;
        border-left: 3px solid var(--accent);
        background: rgba(43,179,177,0.04);
        border-radius: 0 4px 4px 0;
    }

    /* ===== 风险卡片 ===== */
    .risk-card {
        border-radius: var(--radius);
        padding: 18px 22px;
        margin: 12px 0;
        border-left: 5px solid;
    }
    .risk-low { background: #F0FFF4; border-color: var(--risk-low); }
    .risk-med { background: #FFF8F0; border-color: var(--risk-med); }
    .risk-high { background: #FFF4F2; border-color: var(--risk-high); }
    .risk-card-title { font-size: 1.05rem; font-weight: 700; margin-bottom: 8px; }
    .risk-card ul { margin: 6px 0; padding-left: 20px; }
    .risk-card li { font-size: 0.85rem; margin: 3px 0; line-height: 1.6; }

    /* ===== 引导输出 ===== */
    .cl {
        background: var(--bg-card);
        border-left: 3px solid var(--accent);
        border-radius: 0 var(--radius-sm) var(--radius-sm) 0;
        padding: 10px 16px;
        margin: 5px 0;
        font-size: 0.85rem;
        color: var(--text);
        box-shadow: var(--shadow);
    }
    .cl.warn { border-left-color: var(--risk-high); background: #FFF4F2; color: #C0392B; }
    .cl.ok { border-left-color: var(--risk-low); background: #F0FFF4; color: #1E6F3F; }
    .cl.info { border-left-color: var(--accent); background: var(--accent-light); color: #1A7A78; }

    /* ===== 气泡样式 ===== */
    .bubble-user {
        background: var(--primary-light);
        border-radius: 14px 14px 4px 14px;
        padding: 12px 18px;
        margin: 6px 0 6px auto;
        max-width: 70%;
        font-size: 0.88rem;
        line-height: 1.8;
        color: var(--text);
        text-align: right;
    }
    .bubble-ai {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: 14px 14px 14px 4px;
        padding: 12px 18px;
        margin: 6px auto 6px 0;
        max-width: 70%;
        font-size: 0.88rem;
        line-height: 1.8;
        color: var(--text);
        text-align: left;
    }
    .bubble-label { font-size: 0.68rem; color: var(--text-muted); margin-bottom: 4px; }
    .chat-empty { text-align: center; color: var(--text-light); padding: 40px; font-size: 0.88rem; }

    /* ===== 表单 & 输入 ===== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stNumberInput > div > div > input {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-sm) !important;
        color: var(--text) !important;
        font-size: 0.85rem !important;
    }
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus,
    .stNumberInput > div > div > input:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 2px rgba(31,78,140,0.1) !important;
    }
    div[data-testid="stAlert"] { border-radius: var(--radius-sm); }
    .stDataFrame { border-radius: var(--radius-sm); overflow: hidden; }
    .stDateInput > div > div > input {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-sm) !important;
        color: var(--text) !important;
    }
    [data-baseweb="select"] > div {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-sm) !important;
    }

    /* ===== 按钮体系 ===== */
    .stButton > button {
        background: var(--primary) !important;
        color: white !important;
        border: none !important;
        border-radius: var(--radius-sm) !important;
        font-weight: 600 !important;
        font-size: 0.85rem !important;
        padding: 6px 18px !important;
        transition: all 0.2s !important;
    }
    .stButton > button:hover {
        background: var(--primary-dark) !important;
        box-shadow: 0 4px 16px rgba(31,78,140,0.30) !important;
        transform: translateY(-1px);
    }
    .stDownloadButton > button {
        background: var(--accent) !important;
        color: white !important;
        border: none !important;
        border-radius: var(--radius-sm) !important;
        font-weight: 600 !important;
    }
    .stDownloadButton > button:hover {
        background: #259E9C !important;
        box-shadow: 0 4px 14px rgba(43,179,177,0.25) !important;
    }
    /* 次按钮 - 描边 */
    .secondary-btn > .stButton > button {
        background: transparent !important;
        color: var(--primary) !important;
        border: 1.5px solid var(--primary) !important;
    }
    .secondary-btn > .stButton > button:hover {
        background: var(--primary-light) !important;
    }

    /* ===== 指标卡片 ===== */
    [data-testid="metric-container"] {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius) !important;
        padding: 16px 20px !important;
        box-shadow: var(--shadow) !important;
    }
    [data-testid="metric-container"] label {
        color: var(--text-muted) !important;
        font-size: 0.78rem !important;
        font-weight: 500 !important;
    }
    [data-testid="metric-container"] [data-testid="stMetricValue"] {
        color: var(--primary) !important;
        font-weight: 700 !important;
    }
    [data-testid="metric-container"] [data-testid="stMetricDelta"] {
        font-size: 0.75rem !important;
    }

    /* ===== Divider ===== */
    .divider { height: 1px; margin: 20px 0; background: linear-gradient(90deg, transparent, var(--border), transparent); }

    /* ===== Tab 样式 ===== */
    .stTabs [data-baseweb="tab-list"] {
        background: var(--bg-card);
        border-radius: var(--radius);
        padding: 4px;
        gap: 4px;
        border: 1px solid var(--border);
        box-shadow: var(--shadow);
    }
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border-radius: var(--radius-sm);
        color: var(--text-muted);
        padding: 8px 20px;
        font-size: 0.84rem;
    }
    .stTabs [aria-selected="true"] {
        background: var(--primary-light) !important;
        color: var(--primary) !important;
        font-weight: 600;
    }

    /* ===== 进度条 ===== */
    .stProgress > div > div > div { background: var(--accent) !important; }

    /* ===== radio ===== */
    .stRadio > div { gap: 10px; }
    .stRadio [data-baseweb="radio"] span { color: var(--text); }

    /* ===== 材料进度文字 ===== */
    .mat-progress-text { font-size: 0.88rem; font-weight: 600; margin: 6px 0; color: var(--primary); }

    /* ===== 预算面板专用 ===== */
    .budget-section-title {
        font-size: 0.8rem; font-weight: 600; color: var(--text-muted);
        text-transform: uppercase; letter-spacing: 1px;
        margin: 18px 0 12px; padding-bottom: 6px;
        border-bottom: 1px solid var(--border);
    }
    .module-entry-card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius-sm);
        padding: 14px 16px;
        text-align: center;
        cursor: pointer;
        transition: all 0.2s;
        box-shadow: var(--shadow);
    }
    .module-entry-card:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-hover);
        border-color: var(--accent);
    }
    .module-entry-icon { font-size: 1.4rem; margin-bottom: 4px; }
    .module-entry-label { font-size: 0.78rem; color: var(--text); font-weight: 500; }
    .module-entry-desc { font-size: 0.65rem; color: var(--text-muted); margin-top: 2px; }

    /* ===== chat 快捷按钮 ===== */
    .quick-chip {
        background: var(--primary-light);
        border: 1px solid var(--border);
        color: var(--primary);
        border-radius: 20px;
        padding: 4px 14px;
        font-size: 0.75rem;
        cursor: pointer;
        transition: all 0.2s;
        display: inline-block;
        margin: 2px;
    }
    .quick-chip:hover { background: var(--primary); color: white; border-color: var(--primary); }

    /* ===== @keyframes 动画 ===== */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to   { opacity: 1; transform: translateY(0); }
    }
    @keyframes slideUp {
        from { opacity: 0; transform: translateY(20px); }
        to   { opacity: 1; transform: translateY(0); }
    }
    @keyframes fadeInScale {
        from { opacity: 0; transform: scale(0.92); }
        to   { opacity: 1; transform: scale(1); }
    }
    @keyframes fillBar {
        from { width: 0%; }
    }
    @keyframes blink {
        0%, 50% { opacity: 1; }
        51%, 100% { opacity: 0; }
    }

    /* ===== 动画进度条 ===== */
    .anim-progress-wrap {
        width: 100%; height: 8px;
        background: #F0F4F8;
        border-radius: 4px;
        overflow: hidden;
        margin: 4px 0 10px;
    }
    .anim-progress-bar {
        height: 100%;
        border-radius: 4px;
        background: linear-gradient(90deg, #2BB3B1, #1F4E8C);
        animation: fillBar 1s ease-out forwards;
    }
    .anim-progress-bar.warn {
        background: linear-gradient(90deg, #F39C12, #E74C3C);
    }

    /* ===== Hero / 侧边栏入场动画 ===== */
    .hero {
        animation: fadeInScale 0.5s ease-out both;
    }
    .hero-school-badge {
        animation: fadeIn 0.5s ease-out 0.2s both;
    }
    .sidebar-logo {
        animation: fadeIn 0.4s ease-out both;
    }
    .cap-card, .pain-card, .solve-card, .rec-card {
        animation: slideUp 0.45s ease-out both;
    }

    /* ===== 对比表格悬停 ===== */
    .compare-table tbody tr {
        transition: background 0.15s;
    }
    .compare-table tbody tr:hover {
        background: #F0F7FF;
    }

    </style>
    """

st.markdown(get_css(), unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════
# 流程引导类型库 & 用户类型推荐 & 演示案例
# ═══════════════════════════════════════════════════════════════════════
GUIDE_TYPES = {
    "✈️ 差旅费": {
        "required_docs": ["出差审批单（事先填写）", "正规住宿发票", "交通票据（火车票/机票/汽车票）", "打车费收据（如有）", "报销申请表"],
        "tips": "差旅结束后7个工作日内提交报销，超时可能被拒。",
    },
    "🎉 社团/班级活动费": {
        "required_docs": ["活动申请/方案（提前审批）", "参与人员签到名单", "正规发票或收据", "经费使用明细表", "指导老师/辅导员签字"],
        "tips": "活动经费报销须在活动结束后30天内提交。超500元需学院盖章。",
    },
    "🖨️ 办公用品/耗材": {
        "required_docs": ["采购清单（品名、规格、数量、单价）", "正规增值税发票", "验收确认签字"],
        "tips": "单价超500元须填写固定资产登记表，由资产管理处贴条码后方可领用。",
    },
    "🔬 科研经费": {
        "required_docs": ["项目立项批文或合同复印件", "经费支出申请单", "正规发票（注明课题名称）", "项目负责人签字", "设备类须附验收单"],
        "tips": "科研经费须在预算科目内列支，跨科目需提前申请调整，年底前完成结题报销。",
    },
    "🏆 竞赛/培训费": {
        "required_docs": ["竞赛/培训通知原件", "外出审批表（提前填写）", "报名费发票或收据", "参赛/结业证明（返校后）"],
        "tips": "报名费500元以内指导老师审批；500元以上须学院审批。获奖后可申请额外奖励经费。",
    },
    "📦 其他费用": {
        "required_docs": ["费用说明及用途说明", "正规发票或收据", "经手人签字", "部门负责人审批"],
        "tips": "其他费用需详细说明用途，无法提供发票的须填写无票说明，金额较大时需追加审批。",
    },
}

USER_TYPE_RECOMMENDATIONS = {
    "普通学生": {
        "icon": "👤",
        "features": ["💬 智能问答：助学贷款咨询", "💬 智能问答：学费缴纳问题", "🧮 自动计算：竞赛/培训报销"],
    },
    "班委 / 团支书": {
        "icon": "👥",
        "features": ["📋 流程引导：班级活动费报销", "📋 流程引导：团日活动材料清单", "📎 生成清单：Excel 报销清单生成"],
    },
    "社团负责人": {
        "icon": "🌟",
        "features": ["📋 流程引导：社团活动经费", "📋 流程引导：指导老师审批流程", "⚠️ 预算预警：活动预算预警"],
    },
    "科研项目成员": {
        "icon": "🔬",
        "features": ["📋 流程引导：科研经费报销", "📋 流程引导：设备采购审批", "⚠️ 预算预警：预算使用率分析"],
    },
    "教师 / 指导老师": {
        "icon": "👨‍🏫",
        "features": ["📋 流程引导：审批流程查询", "⚠️ 预算预警：项目经费监控", "📎 生成清单：费用清单导出"],
    },
}

DEMO_CASE = {
    "type": "🎉 社团/班级活动费",
    "purpose": "2025级计算机科学与技术5班五四主题团日活动",
    "dept": "数据科学与人工智能学院",
    "applicant": "在校本科生",
    "amount": 680.0,
    "items_desc": "海报打印100元，活动用品180元，奖品300元，矿泉水100元",
    "has_invoice": "✅ 有正规发票",
    "invoice_count": 4,
    "special_note": "本次活动用于五四精神主题团日教育，费用主要用于活动宣传和现场互动。",
    "is_urgent": False,
    "approver": "",
    "submit_loc": "",
}

# ═══════════════════════════════════════════════════════════════════════
# Session State 初始化
# ═══════════════════════════════════════════════════════════════════════
def _init_state():
    defaults = {
        "current_page": "首页总览",
        "chat_history": [],
        "need_ai_reply": False,
        "excel_items": [],
        "guide_data": {
            "type": "✈️ 差旅费",
            "purpose": "",
            "dept": "",
            "applicant": "在校本科生",
            "amount": 0.0,
            "items_desc": "",
            "has_invoice": "✅ 有正规发票",
            "invoice_count": 1,
            "special_note": "",
            "approver": "",
            "submit_loc": "",
            "is_urgent": False,
        },
        "guide_generated": False,
        "prepared_docs": [],
        "missing_docs_count": 0,
        "guide_application_text": "",
        "risk_result": None,
        "user_type": None,
        "demo_loaded": False,
        # 预算结果缓存
        "budget_result": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()

# ═══════════════════════════════════════════════════════════════════════
# AI 调用（带缓存）
# ═══════════════════════════════════════════════════════════════════════
SYSTEM_PROMPT = """你是本校"校园财务问答助手"，帮助学生、教师、社团成员解答财务问题。

【报销标准参考（通用高校惯例，以我校实际规定为准）】
▸ 差旅费
· 教职工住宿：省内≤350元/天，省外≤500元/天
· 学生住宿：≤200元/天
· 餐补：教职工100元/天，学生50元/天
· 交通：凭正规票据实报实销；打车需说明事由
· 火车：原则购二等座；飞机须提前审批

▸ 社团/班级活动
· ≤500元：辅导员/指导老师签字
· 500~2000元：学院审批
· >2000元：学生处/团委审批
· 必需材料：活动方案、参与名单、正规发票

▸ 办公用品/耗材
· 单价≤500元：直接报销（附发票）
· 500~5000元：需填写固定资产登记单
· >5000元：申请政府/学校集中采购

▸ 科研经费
· 区分横向（企业）与纵向（政府/基金会）课题
· 劳务费须签协议并缴税
· 超预算须提前申请调整

▸ 竞赛/培训
· 需提前填写外出审批表
· 报名费凭收据/发票报销，≤500元指导老师审批，>500元学院审批

▸ 助学贷款
· 每年9月办理；流程：申请→学院审核→资助中心审批→银行放款
· 续贷可在线申请，保持良好信用记录

▸ 学校财务基本原则：统一入账、严禁私收费、全面预算管理、年终公开决算

【回答要求】简洁友好、条理清晰。涉及具体金额注明"以我校实际规定为准"。不确定时建议咨询财务处或查阅校园官网。"""

@cache_data(ttl=300)
def call_ai_cached(messages_json: str):
    messages = json.loads(messages_json)
    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        temperature=0.6,
        max_tokens=1024,
    )
    return resp.choices[0].message.content

def call_ai(messages):
    return call_ai_cached(json.dumps(messages, ensure_ascii=False))

def call_ai_stream(messages):
    """流式调用 DeepSeek API，逐 token 返回文本块。"""
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        temperature=0.6,
        max_tokens=1024,
        stream=True,
    )
    for chunk in response:
        delta = chunk.choices[0].delta
        if delta.content:
            yield delta.content

# ═══════════════════════════════════════════════════════════════════════
# 辅助函数（保持不变）
# ═══════════════════════════════════════════════════════════════════════
@cache_data
def parse_reimbursement_cached(text: str):
    d = {}
    m = re.search(r'(\d+(?:\.\d+)?)\s*[天日]', text)
    if m:
        d['days'] = float(m.group(1))
    for pat in [
        r'(?:住宿|酒店|宾馆|旅馆|客栈)[费用]?\s*[：:￥¥]?\s*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*元?[/每]\s*[天日晚夜]',
        r'(\d+(?:\.\d+)?)\s*[元块]/天',
    ]:
        m = re.search(pat, text)
        if m:
            d['hotel_per_day'] = float(m.group(1))
            break
    m = re.search(r'(?:餐[补费]|伙食|用餐|饭补|早餐|午餐|晚餐)[费]?\s*[：:￥¥]?\s*(\d+(?:\.\d+)?)', text)
    if m:
        d['meal_per_day'] = float(m.group(1))
    m = re.search(r'(?:交通|火车|高铁|飞机|机票|大巴|汽车|打车|出租|滴滴)[费票]?\s*[：:￥¥]?\s*(\d+(?:\.\d+)?)', text)
    if m:
        d['transport'] = float(m.group(1))
    m = re.search(r'(?:其他|杂费|额外|补贴)[费]?\s*[：:￥¥]?\s*(\d+(?:\.\d+)?)', text)
    if m:
        d['other'] = float(m.group(1))
    return d if d else None

@cache_data
def parse_items_nl_cached(text: str):
    items = []
    for part in re.split(r'[，,；;\n]+', text):
        part = part.strip()
        if not part:
            continue
        m = re.search(r'^(.+?)\s*[：:￥¥]?\s*(\d+(?:\.\d+)?)\s*[元块]?$', part)
        if m:
            name = m.group(1).strip().strip('：: ')
            price = float(m.group(2))
            if name and price > 0:
                items.append({"物品/费用名称": name, "金额（元）": price})
                continue
        m = re.search(r'^[¥￥]?\s*(\d+(?:\.\d+)?)\s*[元块]?\s+(.+)$', part)
        if m:
            price = float(m.group(1))
            name = m.group(2).strip()
            if name and price > 0:
                items.append({"物品/费用名称": name, "金额（元）": price})
    return items

@cache_data
def build_excel_cached(items_tuple):
    items_list = [{"物品/费用名称": name, "金额（元）": amt} for name, amt in items_tuple]
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook()
    ws = wb.active
    ws.title = "报销清单"
    thin = Side(style='thin', color="CCCCCC")
    bdr = Border(left=thin, right=thin, top=thin, bottom=thin)
    gold_fill = PatternFill("solid", fgColor="D4AF37")
    navy_fill = PatternFill("solid", fgColor="1E3A5F")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    gray_fill = PatternFill("solid", fgColor="EBF0FA")
    ws.merge_cells('A1:E1')
    ws['A1'] = "报 销 费 用 清 单"
    ws['A1'].font = Font(name='微软雅黑', size=17, bold=True, color="1E3A5F")
    ws['A1'].fill = gold_fill
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 40
    ws.merge_cells('A2:E2')
    ws['A2'] = f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
    ws['A2'].font = Font(name='微软雅黑', size=9, color="888888")
    ws['A2'].alignment = Alignment(horizontal='right', vertical='center')
    ws.row_dimensions[2].height = 18
    headers = ["序号", "物品 / 费用名称", "单价（元）", "数量", "金额（元）"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(3, col, h)
        c.fill = navy_fill
        c.font = Font(name='微软雅黑', bold=True, color="FFFFFF", size=11)
        c.alignment = Alignment(horizontal='center', vertical='center')
        c.border = bdr
    ws.row_dimensions[3].height = 28
    for i, item in enumerate(items_list):
        r = i + 4
        vals = [i + 1, item["物品/费用名称"], item["金额（元）"], 1, item["金额（元）"]]
        fill = white_fill if i % 2 == 0 else gray_fill
        for col, val in enumerate(vals, 1):
            c = ws.cell(r, col, val)
            c.fill = fill
            c.font = Font(name='微软雅黑', size=10)
            c.border = bdr
            c.alignment = Alignment(horizontal='center' if col in [1,3,4,5] else 'left', vertical='center')
        ws.row_dimensions[r].height = 22
    tr = len(items_list) + 4
    total = sum(x["金额（元）"] for x in items_list)
    ws.merge_cells(f'A{tr}:D{tr}')
    ws[f'A{tr}'] = "合　　计"
    ws[f'A{tr}'].font = Font(name='微软雅黑', bold=True, color="FFFFFF", size=12)
    ws[f'A{tr}'].fill = navy_fill
    ws[f'A{tr}'].alignment = Alignment(horizontal='center', vertical='center')
    ws[f'E{tr}'] = total
    ws[f'E{tr}'].font = Font(name='微软雅黑', bold=True, color="D4AF37", size=13)
    ws[f'E{tr}'].fill = navy_fill
    ws[f'E{tr}'].alignment = Alignment(horizontal='center', vertical='center')
    for col in range(1, 6):
        ws.cell(tr, col).border = bdr
    ws.row_dimensions[tr].height = 32
    sr = tr + 1
    ws.merge_cells(f'A{sr}:E{sr}')
    ws[f'A{sr}'] = f"报销总金额：¥ {total:.2f} 元"
    ws[f'A{sr}'].font = Font(name='微软雅黑', size=10, color="555555", italic=True)
    ws[f'A{sr}'].alignment = Alignment(horizontal='right', vertical='center')
    ws.row_dimensions[sr].height = 22
    sign_r = tr + 3
    ws.merge_cells(f'A{sign_r}:E{sign_r}')
    ws[f'A{sign_r}'] = "报销人：_______________ 审批人：_______________ 财务审核：_______________ 日期：_______________"
    ws[f'A{sign_r}'].font = Font(name='微软雅黑', size=10, color="444444")
    ws[f'A{sign_r}'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[sign_r].height = 32
    for col, w in [('A',7),('B',30),('C',14),('D',8),('E',14)]:
        ws.column_dimensions[col].width = w
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue(), total

def parse_reimbursement(text):
    return parse_reimbursement_cached(text)

def parse_items_nl(text):
    return parse_items_nl_cached(text)

def build_excel(items_list):
    items_tuple = tuple((item["物品/费用名称"], item["金额（元）"]) for item in items_list)
    return build_excel_cached(items_tuple)

def get_auto_approver_submit(gtype, amount):
    if "差旅" in gtype:
        if amount <= 500:
            return "指导老师", "学院办公室"
        else:
            return "学院分管院长", "财务处报销大厅"
    elif "社团" in gtype:
        if amount <= 500:
            return "社团指导老师", "团委财务窗口"
        else:
            return "学院学工办", "团委财务窗口"
    elif "办公" in gtype:
        return "实验室/部门负责人", "资产与实验室管理处"
    elif "科研" in gtype:
        return "项目负责人", "科研院财务科"
    elif "竞赛" in gtype:
        return "竞赛指导老师", "教务处实践科"
    else:
        return "部门负责人", "财务处"

def generate_guide_output(data):
    gtype = data["type"]
    type_info = GUIDE_TYPES.get(gtype, GUIDE_TYPES["📦 其他费用"])
    amount = data["amount"]
    has_inv = data["has_invoice"]
    applicant = data["applicant"]
    docs_html = "".join([f'<div class="cl">📄 {doc}</div>' for doc in type_info["required_docs"]])
    warn_html = ""
    if amount > 2000:
        warn_html += '<div class="cl warn">🚨 金额超过2000元，需要学院/学生处级别审批，请提前联系。</div>'
    elif amount > 500:
        warn_html += '<div class="cl info">ℹ️ 金额超过500元，需学院盖章或中级审批，请确认审批流程。</div>'
    else:
        warn_html += '<div class="cl ok">✅ 金额在500元以内，指导老师/辅导员签字即可。</div>'
    if "❌" in has_inv:
        warn_html += '<div class="cl warn">🚨 无票据：请联系商家补开发票，或提交书面无票说明（需部门负责人签字）。</div>'
    elif "⚠️" in has_inv:
        warn_html += '<div class="cl info">ℹ️ 仅有收据：部分情况可接受，建议提前与财务处确认是否可行。</div>'
    else:
        warn_html += '<div class="cl ok">✅ 持有正规发票，票据符合基本要求。</div>'
    if "学生" in applicant and "差旅" in gtype:
        warn_html += '<div class="cl info">ℹ️ 学生住宿标准≤200元/天，超出部分需额外说明。</div>'
    if data.get("special_note"):
        warn_html += f'<div class="cl info">📝 备注：{data["special_note"]}</div>'
    if data.get("is_urgent"):
        warn_html += '<div class="cl warn">⚡ 已标记紧急报销，提交时请在报销单封面注明「加急」并说明原因。</div>'

    approver = data.get("approver")
    submit_loc = data.get("submit_loc")
    if not approver:
        approver, _ = get_auto_approver_submit(gtype, amount)
    if not submit_loc:
        _, submit_loc = get_auto_approver_submit(gtype, amount)

    steps_html = ""
    steps_list = [
        f"整理全部票据并按顺序粘贴在报销单背面",
        f"填写报销申请表（学院财务/行政办公室领取）",
        f"提交给审批人：**{approver}** 签字确认",
        f"携带全部材料前往：**{submit_loc}** 提交",
        f"留存报销单复印件或拍照备份，等待到账（一般3-10个工作日）",
    ]
    for i, s in enumerate(steps_list, 1):
        steps_html += f'<div class="cl">{i}. {s}</div>'
    tips_html = f'<div class="cl info">💡 **小贴士：** {type_info["tips"]}</div>'
    contact_html = '<div class="cl info">📞 如有疑问，请拨打校财务处咨询电话，或前往窗口现场确认。</div>'
    return docs_html + warn_html + steps_html + tips_html + contact_html

def generate_application_text(data):
    purpose = data.get("purpose", "").strip() or "未填写事由"
    amount = data.get("amount", 0)
    has_invoice = data.get("has_invoice", "")
    items_desc = data.get("items_desc", "").strip()
    special_note = data.get("special_note", "").strip()
    dept = data.get("dept", "").strip() or "未填写学院"

    text = f"本次报销事项为 {purpose}，"
    if items_desc:
        text += f"主要费用包括{items_desc}，"
    text += f"共计 {amount:.0f} 元。"
    if "正规发票" in has_invoice:
        text += "相关票据齐全"
    elif "只有收据" in has_invoice:
        text += "仅有收据作为凭证"
    else:
        text += "暂无正规票据"
    text += "，费用用途明确，符合学校财务相关规定。"
    if special_note:
        text += f"特此说明：{special_note}。"
    if dept and dept != "未填写学院":
        text += f"\n申请人所属{dept}，"
    text += "现申请办理报销，请老师审核。"
    return text

def calculate_reimbursement_risk(data, missing_docs_count=0):
    amount = data.get("amount", 0)
    has_invoice = data.get("has_invoice", "")
    is_urgent = data.get("is_urgent", False)

    risk_score = 0
    reasons = []
    suggestions = []

    if amount > 500:
        risk_score += 1
        reasons.append(f"报销金额 {amount:.0f} 元超过 500 元，需要学院审批")
        suggestions.append("提前联系学院办公室确认审批要求")
    if amount > 2000:
        risk_score += 2
        reasons.append(f"报销金额 {amount:.0f} 元超过 2000 元，需学院/学生处级别审批")
        suggestions.append("提前联系财务处确认审批流程")
    if "只有收据" in has_invoice:
        risk_score += 1
        reasons.append("仅有收据，部分情况可能无法全额报销")
        suggestions.append("提前与财务处确认收据是否可行")
    if "无任何票据" in has_invoice:
        risk_score += 3
        reasons.append("无任何票据，报销申请可能被退回")
        suggestions.append("联系商家补开发票或准备书面无票说明")
    if is_urgent:
        risk_score += 1
        reasons.append("标记为紧急报销，需加急处理")
        suggestions.append("在报销单封面注明「加急」并说明原因")
    if missing_docs_count > 0:
        risk_score += missing_docs_count
        reasons.append(f"当前仍缺少 {missing_docs_count} 项材料")
        suggestions.append("补齐缺失材料后再提交")

    if risk_score <= 1:
        level = "低风险"
        default_suggestions = ["当前报销风险较低，请按流程正常提交"]
    elif risk_score <= 3:
        level = "中风险"
        default_suggestions = ["建议提交前再次确认材料完整性"]
    else:
        level = "高风险"
        default_suggestions = ["建议暂停提交，先补齐所有材料并咨询财务处"]

    if not suggestions:
        suggestions = default_suggestions

    summary_reason = "；".join(reasons[:3]) if reasons else "未发现明显风险因素"

    return {
        "level": level,
        "score": risk_score,
        "reasons": reasons if reasons else ["未发现明显风险因素"],
        "suggestions": suggestions,
        "summary": summary_reason,
    }

def check_missing_docs(required_docs, prepared_docs):
    prepared = [d for d in required_docs if d in prepared_docs]
    missing = [d for d in required_docs if d not in prepared_docs]
    return prepared, missing

def build_reimbursement_docx(data, application_text, risk_result, prepared_docs, missing_docs):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("报销申请说明")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x2C, 0x3C)

    doc.add_paragraph()

    doc.add_paragraph("一、基本信息", style='Heading 2')
    info_items = [
        ("报销类型", data.get("type", "")),
        ("报销事由", data.get("purpose", "")),
        ("所在学院", data.get("dept", "")),
        ("申请人身份", data.get("applicant", "")),
        ("报销金额", f"{data.get('amount', 0):.0f} 元"),
        ("票据情况", data.get("has_invoice", "")),
        ("发票张数", str(data.get("invoice_count", 0))),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}：")
        run_label.bold = True
        run_label.font.size = Pt(10.5)
        run_value = p.add_run(value)
        run_value.font.size = Pt(10.5)

    if data.get("items_desc", "").strip():
        doc.add_paragraph("二、费用说明", style='Heading 2')
        doc.add_paragraph(data["items_desc"].strip())

    doc.add_paragraph("三、材料准备情况", style='Heading 2')
    for d in prepared_docs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"✅ {d}").font.size = Pt(10.5)
    for d in missing_docs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"❌ {d}（未准备）").font.size = Pt(10.5)

    doc.add_paragraph("四、报销风险等级", style='Heading 2')
    level_map = {"低风险": "低风险", "中风险": "中风险", "高风险": "高风险"}
    p = doc.add_paragraph()
    run = p.add_run(f"风险等级：{level_map.get(risk_result['level'], risk_result['level'])}")
    run.bold = True
    run.font.size = Pt(11)
    for reason in risk_result["reasons"]:
        doc.add_paragraph(reason, style='List Bullet')

    doc.add_paragraph("五、报销申请说明", style='Heading 2')
    doc.add_paragraph(application_text)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def load_demo_case():
    for k, v in DEMO_CASE.items():
        st.session_state.guide_data[k] = v
    st.session_state.guide_generated = True
    st.session_state.prepared_docs = []
    st.session_state.missing_docs_count = 0
    st.session_state.guide_application_text = ""
    st.session_state.risk_result = None

def show_risk_card(risk_result):
    level = risk_result["level"]
    if level == "低风险":
        css_class = "risk-low"
        icon = "✅"
        color_label = "低风险"
    elif level == "中风险":
        css_class = "risk-med"
        icon = "⚠️"
        color_label = "中风险"
    else:
        css_class = "risk-high"
        icon = "🚨"
        color_label = "高风险"

    bg_colors = {"低风险": "#F0FFF4", "中风险": "#FFF8F0", "高风险": "#FFF4F2"}
    border_colors = {"低风险": "#27AE60", "中风险": "#F39C12", "高风险": "#E74C3C"}

    reasons_html = "".join([f"<li>{r}</li>" for r in risk_result["reasons"]])
    suggestions_html = "".join([f"<li>💡 {s}</li>" for s in risk_result["suggestions"]])

    card_html = f"""
    <div style="background:{bg_colors[level]};border-left:5px solid {border_colors[level]};border-radius:14px;padding:20px 24px;margin:12px 0;">
        <div style="font-size:1.1rem;font-weight:700;margin-bottom:8px;">{icon} 当前报销风险等级：{color_label}</div>
        <div style="margin:10px 0;">
            <div style="font-weight:600;font-size:0.9rem;margin-bottom:4px;">原因：</div>
            <ul style="margin:4px 0;padding-left:20px;">{reasons_html}</ul>
        </div>
        <div>
            <div style="font-weight:600;font-size:0.9rem;margin-bottom:4px;">建议：</div>
            <ul style="margin:4px 0;padding-left:20px;">{suggestions_html}</ul>
        </div>
    </div>
    """
    st.markdown(card_html, unsafe_allow_html=True)

def render_animated_kpi_cards():
    """首页 KPI 卡片 —— 使用 JS 计数器动画。"""
    cards_data = [
        {"icon": "📂", "value": 6, "suffix": " 类报销场景", "desc": "覆盖差旅、活动、办公、科研、竞赛、其他"},
        {"icon": "🗣️", "value": 100, "suffix": "% 自然语言", "desc": "用日常用语描述费用，AI 自动理解"},
        {"icon": "📥", "value": 1, "suffix": " 键导出 Excel", "desc": "标准报销清单格式，即生即下"},
        {"icon": "📊", "value": 3, "suffix": " 级风险预警", "desc": "智能识别材料缺失与预算风险"},
    ]
    cards_json = json.dumps(cards_data, ensure_ascii=False)

    components.html(f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
body {{ margin:0; padding:0; background:transparent; font-family:'Noto Sans SC',-apple-system,sans-serif; }}
.grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:14px; }}
.card {{ background:#FFFFFF; border:1px solid #E5EEF3; border-radius:12px; padding:20px 16px;
  text-align:center; box-shadow:0 2px 12px rgba(31,78,140,0.06); transition:transform 0.2s,box-shadow 0.2s; }}
.card:hover {{ transform:translateY(-3px); box-shadow:0 6px 24px rgba(31,78,140,0.10); }}
.icon {{ font-size:1.6rem; margin-bottom:8px; }}
.value {{ font-size:1.1rem; font-weight:700; color:#1F4E8C; }}
.label {{ font-size:0.78rem; color:#7A8FA0; margin-top:2px; }}
</style></head><body>
<div class="grid" id="kpi-grid"></div>
<script>
(function(){{
  var cards={cards_json};
  var html='';
  cards.forEach(function(c){{
    html+='<div class="card"><div class="icon">'+c.icon+'</div>';
    html+='<div class="value"><span class="count" data-target="'+c.value+'">0</span>'+c.suffix+'</div>';
    html+='<div class="label">'+c.desc+'</div></div>';
  }});
  document.getElementById('kpi-grid').innerHTML=html;
  var counters=document.querySelectorAll('.count');
  counters.forEach(function(el){{
    var target=parseInt(el.dataset.target), current=0, step=target>50?Math.ceil(target/60):1;
    var timer=setInterval(function(){{
      current+=step; if(current>=target){{ current=target; clearInterval(timer); }}
      el.textContent=current;
    }},25);
  }});
}})();
</script></body></html>""", height=150, scrolling=False)

def render_typewriter_message(text: str, key_suffix: str = ""):
    """使用 JS 打字机动画渲染 AI 回复。"""
    escaped = json.dumps(text, ensure_ascii=False)
    lines = text.count('\n') + 1
    height = max(80, lines * 26 + 50)

    components.html(f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
body {{ margin:0; padding:0; background:transparent; font-family:'Noto Sans SC',-apple-system,sans-serif; }}
.bubble-ai {{
  background:#FFFFFF; border:1px solid #E5EEF3; border-radius:14px 14px 14px 4px;
  padding:12px 18px; font-size:0.88rem; line-height:1.8; color:#1A2C3C;
  white-space:pre-wrap; word-break:break-word;
}}
.cursor {{ display:inline-block; width:2px; height:1em; background:#1F4E8C;
  animation:blink 0.8s infinite; vertical-align:text-bottom; margin-left:2px; }}
@keyframes blink {{ 0%,50% {{ opacity:1; }} 51%,100% {{ opacity:0; }} }}
</style></head><body>
<div class="bubble-ai"><span id="tw{key_suffix}"></span><span class="cursor" id="cur{key_suffix}"></span></div>
<script>
(function(){{
  var t={escaped};
  var i=0, el=document.getElementById("tw{key_suffix}"), cur=document.getElementById("cur{key_suffix}");
  function type(){{
    if(i<t.length){{ el.textContent+=t.charAt(i); i++; setTimeout(type,18+Math.random()*12); }}
    else {{ cur.style.display='none'; }}
  }}
  type();
}})();
</script></body></html>""", height=height, scrolling=False)

# ═══════════════════════════════════════════════════════════════════════
# 左侧固定导航栏
# ═══════════════════════════════════════════════════════════════════════
NAV_ITEMS = ["🏠 首页总览", "💬 智能问答", "🧮 自动计算", "📎 生成清单", "📋 流程引导", "⚠️ 预算预警"]

with st.sidebar:
    # Logo & 品牌区
    badge_img = get_school_badge()
    img_tag = f'<img src="{badge_img}" alt="校徽">' if badge_img else '<span>江财</span>'
    st.markdown(f"""
    <div class="sidebar-logo">
        <div class="sidebar-badge">
            {img_tag}
        </div>
        <div class="sidebar-sysname">校园财务问答助手</div>
        <div class="sidebar-subtitle">智慧财务 · 高效合规</div>
        <div class="sidebar-school">江西财经大学现代经济管理学院</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="nav-section-label">功能导航</div>', unsafe_allow_html=True)

    # 导航菜单 - 使用 radio 实现
    selected = st.radio(
        "导航",
        NAV_ITEMS,
        index=NAV_ITEMS.index(st.session_state.current_page) if st.session_state.current_page in NAV_ITEMS else 0,
        label_visibility="collapsed",
        key="nav_radio",
    )

    st.session_state.current_page = selected

    # 底部用户信息区
    st.markdown("""
    <div style="text-align:center;padding:16px 16px 12px;margin-top:24px;border-top:1px solid #E5EEF3;">
        <div style="font-size:2rem;margin-bottom:6px;">👤</div>
        <div style="font-weight:600;font-size:0.82rem;color:#1A2C3C;margin-bottom:2px;">郭志明</div>
        <div style="font-size:0.65rem;color:#7A8FA0;line-height:1.5;">
            数据科学与人工智能学院<br>
            2025级 · 本科生
        </div>
        <div style="margin-top:8px;font-size:0.58rem;color:#A0B4C4;">
            v2.1 · 比赛演示版
        </div>
    </div>
    """, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════
# 主内容区 - 页面路由
# ═══════════════════════════════════════════════════════════════════════
page = st.session_state.current_page

# ============================================================
# 页面：首页总览
# ============================================================
if page == "🏠 首页总览":
    # Hero 区
    col_h_left, col_h_right = st.columns([3, 1.2])
    with col_h_left:
        st.markdown("""
        <div class="hero">
            <div class="hero-content">
                <div class="hero-title">校园财务问答助手</div>
                <div class="hero-sub">AI问答 · 报销计算 · 清单生成 · 流程引导 · 预算预警</div>
                <div>
                    <span class="hero-badge">💡 懂财务</span>
                    <span class="hero-badge">🤖 更智能</span>
                    <span class="hero-badge">⚡ 更高效</span>
                    <span class="hero-badge">✅ 更合规</span>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with col_h_right:
        badge_img = get_school_badge()
        badge_tag = f'<img src="{badge_img}" alt="校徽">' if badge_img else '江财'
        st.markdown(f"""
        <div style="display:flex;align-items:center;justify-content:flex-end;height:100%;padding-right:12px;">
            <div class="hero-school-badge">
                {badge_tag}
            </div>
        </div>
        """, unsafe_allow_html=True)

    # 四个能力标签卡（带 JS 计数器动画）
    render_animated_kpi_cards()

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # 痛点 vs 方案
    st.markdown("#### 🔍 校园报销痛点 vs 💡 我们的方案")
    p_col1, p_col2 = st.columns(2)
    pains = [
        ("📋", "报销政策分散", "学生不知道能不能报、能报多少，翻文件效率低。"),
        ("🔄", "报销流程复杂", "不清楚找谁审批、去哪里提交，环节多易出错。"),
        ("📑", "材料容易遗漏", "凭经验准备材料，缺一件就被退回。"),
        ("🧮", "金额计算易错", "住宿、餐补、交通…手动加总易出错。"),
        ("⚠️", "预算风险难判断", "临近截止才发现超支或快过期。"),
    ]
    solutions = [
        ("💬", "AI 智能问答", "7×24 在线回答财务政策问题。"),
        ("🧮", "自动计算金额", "根据输入自动估算，识别超标项。"),
        ("📎", "Excel 清单生成", "标准格式报销清单，一键下载。"),
        ("📋", "流程与材料引导", "自动生成材料清单、审批流程。"),
        ("⚠️", "预算预警与分析", "可视化预算使用率，提前规避超支。"),
    ]
    with p_col1:
        st.markdown("<div style='font-weight:600;color:#E74C3C;margin-bottom:8px;'>❌ 传统痛点</div>", unsafe_allow_html=True)
        for icon, title, desc in pains:
            st.markdown(f'<div class="pain-card"><div style="font-weight:600;font-size:0.9rem;">{icon} {title}</div><div style="font-size:0.78rem;color:#7A8FA0;">{desc}</div></div>', unsafe_allow_html=True)
    with p_col2:
        st.markdown("<div style='font-weight:600;color:var(--accent);margin-bottom:8px;'>✅ 本系统方案</div>", unsafe_allow_html=True)
        for icon, title, desc in solutions:
            st.markdown(f'<div class="solve-card"><div style="font-weight:600;font-size:0.9rem;">{icon} {title}</div><div style="font-size:0.78rem;color:#7A8FA0;">{desc}</div></div>', unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # 用户类型推荐
    @st.fragment
    def home_user_type_fragment():
        st.markdown("#### 👤 选择你的身份，获取推荐功能")
        user_type_options = list(USER_TYPE_RECOMMENDATIONS.keys())
        current_type = st.session_state.get("user_type")
        default_idx = 0
        if current_type and current_type in user_type_options:
            default_idx = user_type_options.index(current_type)
        selected_type = st.radio(
            "我是：",
            user_type_options,
            horizontal=True,
            index=default_idx,
            key="home_user_type",
            label_visibility="collapsed",
        )
        st.session_state.user_type = selected_type

        if selected_type:
            rec = USER_TYPE_RECOMMENDATIONS[selected_type]
            st.markdown(f'<div class="rec-card"><div class="rec-title">{rec["icon"]} {selected_type} · 推荐功能</div></div>', unsafe_allow_html=True)
            for feat in rec["features"]:
                st.markdown(f'<div class="rec-item">{feat}</div>', unsafe_allow_html=True)
    home_user_type_fragment()

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # 传统 vs AI 对比
    st.markdown("#### ⚡ 传统方式 vs 本系统")
    compare_html = """
    <table class="compare-table">
        <tr><th style="width:22%;">对比项</th><th style="width:39%;">传统方式</th><th style="width:39%;">本系统</th></tr>
        <tr><td>查询政策</td><td>翻文件、问老师、等回复</td><td>✅ AI 直接问答，秒级响应</td></tr>
        <tr><td>计算金额</td><td>手动计算，容易出错</td><td>✅ 自动识别并计算，即时校验</td></tr>
        <tr><td>准备材料</td><td>靠经验，容易遗漏</td><td>✅ 自动生成清单并检查缺失</td></tr>
        <tr><td>审批流程</td><td>不知道找谁签字</td><td>✅ 自动提示审批人和提交地点</td></tr>
        <tr><td>预算风险</td><td>事后发现超支</td><td>✅ 提前预警和分析</td></tr>
    </table>
    """
    st.markdown(compare_html, unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # 报销类型分布饼图
    @st.fragment
    def home_pie_fragment():
        st.markdown("#### 📊 常见报销类型分布")
        st.caption("基于校园实际报销数据的模拟分布")
        pie_data = {
            "差旅费": 35,
            "社团/班级活动费": 28,
            "办公用品/耗材": 18,
            "科研经费": 12,
            "竞赛/培训费": 7,
        }
        colors_pie = ["#1F4E8C", "#2BB3B1", "#3498DB", "#F39C12", "#E74C3C"]
        fig_pie = go.Figure(go.Pie(
            labels=list(pie_data.keys()),
            values=list(pie_data.values()),
            hole=0.4,
            marker=dict(colors=colors_pie, line=dict(color="#FFFFFF", width=2)),
            textinfo="label+percent",
            textfont=dict(family="Noto Sans SC", size=11, color="#1A2C3C"),
            hovertemplate="%{label}: %{value}%<extra></extra>",
        ))
        fig_pie.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(family="Noto Sans SC"),
            height=320,
            margin=dict(t=10, b=10, l=10, r=10),
            showlegend=False,
        )
        st.plotly_chart(fig_pie, use_container_width=True)
    home_pie_fragment()

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # 演示案例
    @st.fragment
    def home_demo_fragment():
        st.markdown("#### 🎯 快速体验完整功能")
        st.caption("点击下方按钮，自动填入演示案例数据，然后切换到「📋 流程引导」页面查看完整闭环。")
        if st.button("🚀 一键加载演示案例", key="demo_case_home", use_container_width=True):
            load_demo_case()
            st.success("✅ 演示案例已加载！请切换到「📋 流程引导」页面查看材料清单、风险等级和报销说明。")
    home_demo_fragment()

# ============================================================
# 页面：智能问答
# ============================================================
elif page == "💬 智能问答":
    st.markdown('<div class="card"><div class="card-title">💬 智能财务问答</div>', unsafe_allow_html=True)
    st.caption("支持多轮连续对话，所有回答基于学校财务制度及通用高校惯例")
    QUICK_QS = [
        "差旅费住宿标准是多少？", "社团活动报销需要什么材料？",
        "助学贷款怎么办理？", "科研经费报销设备的流程？",
        "报销单丢失怎么办？", "办公用品超过500元怎么处理？",
        "打车费可以报销吗？", "横向课题和纵向课题有什么区别？",
    ]

    def add_quick_question(q):
        st.session_state.chat_history.append({"role": "user", "content": q})
        st.session_state.need_ai_reply = True

    def send_question():
        user_q = st.session_state.qa_input
        if user_q.strip():
            st.session_state.chat_history.append({"role": "user", "content": user_q.strip()})
            st.session_state.need_ai_reply = True

    def clear_chat():
        st.session_state.chat_history = []
        st.session_state.need_ai_reply = False

    @st.fragment
    def qa_interactive():
        st.markdown("**💡 快捷提问：**")
        cols = st.columns(4)
        for i, q in enumerate(QUICK_QS):
            with cols[i % 4]:
                st.button(q, key=f"quick_{i}", on_click=add_quick_question, args=(q,), use_container_width=True)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        if st.session_state.get("need_ai_reply"):
            st.session_state.need_ai_reply = False
            with st.spinner("思考中..."):
                msgs = [{"role": "system", "content": SYSTEM_PROMPT}] + st.session_state.chat_history
                reply = call_ai(msgs)
                st.session_state.chat_history.append({"role": "assistant", "content": reply})

        if not st.session_state.chat_history:
            st.markdown('<div class="chat-empty">👆 点击快捷提问或在下方输入框开始对话</div>', unsafe_allow_html=True)
        else:
            total = len(st.session_state.chat_history)
            for idx, msg in enumerate(st.session_state.chat_history):
                if msg["role"] == "user":
                    st.markdown(f'<div class="bubble-label user-label" style="text-align:right;">👤 你</div>'
                                f'<div class="bubble-user">{msg["content"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="bubble-label ai-label">🤖 财务助手</div>', unsafe_allow_html=True)
                    is_latest = (idx == total - 1)
                    if is_latest and len(msg["content"]) <= 600:
                        render_typewriter_message(msg["content"], key_suffix=f"_{idx}")
                    else:
                        st.markdown(f'<div class="bubble-ai">{msg["content"]}</div>', unsafe_allow_html=True)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns([6, 1, 1])
        with c1:
            st.text_input("输入问题", key="qa_input", label_visibility="collapsed",
                          placeholder="例：差旅报销需要几天内提交？餐补标准是多少？")
        with c2:
            st.button("发送 →", key="qa_send", on_click=send_question, use_container_width=True)
        with c3:
            st.button("清空", key="qa_clear", on_click=clear_chat, use_container_width=True)

    qa_interactive()

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# 页面：自动计算
# ============================================================
elif page == "🧮 自动计算":
    st.markdown('<div class="card"><div class="card-title">🧮 自动计算报销金额</div>', unsafe_allow_html=True)

    @st.fragment
    def calc_fragment():
        mode_calc = st.radio("输入方式", ["📝 自然语言描述", "🔢 手动填写表单"], horizontal=True, key="calc_mode")
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        if mode_calc == "📝 自然语言描述":
            st.caption("💡 直接描述出差情况，支持多种表达方式")
            examples = [
                "出差3天，住宿400元/天，餐补100元/天，交通费200元",
                "去北京开会5天，酒店350元每晚，伙食补贴80元/天，高铁来回共600元",
                "差旅2日，宾馆每天280，打车费共150",
            ]
            st.markdown(f"**参考格式：** `{examples[0]}` · `{examples[1]}`")
            nl_input = st.text_area("描述出差/报销情况", height=100, key="calc_nl", placeholder=examples[0])
            if st.button("🧮 立即计算", key="calc_nl_btn"):
                result = parse_reimbursement(nl_input)
                if not result:
                    st.warning("⚠️ 未能识别有效数据。请包含「天数」「住宿」「餐补」「交通」等关键词。")
                else:
                    days = result.get('days', 0)
                    hotel = result.get('hotel_per_day', 0)
                    meal = result.get('meal_per_day', 0)
                    trans = result.get('transport', 0)
                    other = result.get('other', 0)
                    hotel_t = days * hotel
                    meal_t = days * meal
                    total = hotel_t + meal_t + trans + other
                    st.success(f"✅ 预计可报销总额：**{total:.2f} 元**")
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("出差天数", f"{days} 天")
                    c2.metric("住宿小计", f"{hotel_t:.0f} 元")
                    c3.metric("餐补小计", f"{meal_t:.0f} 元")
                    c4.metric("交通 + 其他", f"{trans + other:.0f} 元")
                    st.caption("*本结果仅供参考，最终以学校财务处核定为准。*")
        else:
            @st.fragment
            def manual_calc_fragment():
                st.caption("逐项填写各项费用，自动汇总计算")
                c1, c2 = st.columns(2)
                with c1:
                    days = st.number_input("出差天数（天）", min_value=0, step=1, key="m_days_frag")
                    hotel_day = st.number_input("住宿费（元/天）", min_value=0.0, step=50.0, key="m_hotel_frag")
                    meal_day = st.number_input("餐补（元/天）", min_value=0.0, step=10.0, key="m_meal_frag")
                with c2:
                    transport = st.number_input("交通费（元，合计）", min_value=0.0, step=10.0, key="m_trans_frag")
                    other = st.number_input("其他费用（元）", min_value=0.0, step=10.0, key="m_other_frag")
                    person = st.selectbox("申请人身份", ["在校学生", "教职工", "研究生"], key="m_person_frag")
                if st.button("🧮 计算合计", key="calc_manual_frag"):
                    hotel_t = days * hotel_day
                    meal_t = days * meal_day
                    total = hotel_t + meal_t + transport + other
                    if person == "在校学生" and hotel_day > 200:
                        st.warning(f"⚠️ 住宿费 {hotel_day:.0f} 元/天超过学生标准（≤200元），超出部分可能无法全额报销。")
                    if person == "教职工" and hotel_day > 350:
                        st.warning(f"⚠️ 住宿费 {hotel_day:.0f} 元/天超过省内教职工标准（≤350元）。")
                    st.success(f"✅ 预计可报销总额：**{total:.2f} 元**")
                    col1, col2, col3, col4, col5 = st.columns(5)
                    col1.metric("天数", f"{days}天")
                    col2.metric("住宿小计", f"{hotel_t:.0f} 元")
                    col3.metric("餐补小计", f"{meal_t:.0f} 元")
                    col4.metric("交通费", f"{transport:.0f} 元")
                    col5.metric("汇总总额", f"{total:.2f} 元")
                    if total > 0:
                        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
                        st.markdown("**费用结构：**")
                        for label, val in [("住宿", hotel_t), ("餐补", meal_t), ("交通", transport), ("其他", other)]:
                            if val > 0:
                                pct = val / total
                                st.write(f"{label} {val:.2f} 元 （{pct*100:.1f}%）")
                                st.markdown(f'<div class="anim-progress-wrap"><div class="anim-progress-bar" style="width:{pct*100:.1f}%;animation:fillBar 0.8s ease-out forwards;"></div></div>', unsafe_allow_html=True)
                        st.caption("*本结果仅供参考，最终以学校财务处核定为准。*")
            manual_calc_fragment()

    calc_fragment()

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# 页面：生成清单
# ============================================================
elif page == "📎 生成清单":
    st.markdown('<div class="card"><div class="card-title">📎 生成报销清单 Excel</div>', unsafe_allow_html=True)

    def delete_excel_item(index):
        st.session_state.excel_items.pop(index)

    def clear_excel_items():
        st.session_state.excel_items = []

    @st.fragment
    def list_fragment():
        mode_xl = st.radio("录入方式", ["📝 自然语言批量输入", "➕ 逐条手动添加"], horizontal=True, key="xl_mode")
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        if mode_xl == "📝 自然语言批量输入":
            st.caption("💡 用逗号或换行分隔每一项，格式非常灵活")
            examples_xl = "打印费50元，笔记本30，马克笔20元\n会场布置费200，茶水服务费80，海报印刷60元"
            nl_xl = st.text_area("输入费用明细", value="", height=140, key="xl_nl", placeholder=examples_xl)
            if st.button("📊 解析并预览", key="xl_parse_btn"):
                items = parse_items_nl(nl_xl)
                if not items:
                    st.warning("⚠️ 未能识别费用条目。请确保每条包含「名称」和「金额」，如：打印费50元")
                else:
                    st.session_state.excel_items = items
                    st.success(f"✅ 成功识别 {len(items)} 条费用")
        else:
            @st.fragment
            def manual_add_fragment():
                st.caption("依次输入每项费用名称和金额，支持随时删除")
                ca, cb, cc = st.columns([3, 2, 1])
                with ca:
                    new_name = st.text_input("费用名称", key="xl_name_frag", placeholder="例：打印费")
                with cb:
                    new_amt = st.number_input("金额（元）", min_value=0.0, step=1.0, key="xl_amt_frag")
                with cc:
                    st.write("")
                    st.write("")
                    if st.button("➕ 添加", key="xl_add_btn_frag", use_container_width=True):
                        if new_name.strip() and new_amt > 0:
                            st.session_state.excel_items.append({"物品/费用名称": new_name.strip(), "金额（元）": new_amt})
            manual_add_fragment()

        if st.session_state.excel_items:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.markdown(f"**📋 当前清单（共 {len(st.session_state.excel_items)} 条）**")
            for i, item in enumerate(st.session_state.excel_items):
                col_n, col_p, col_d = st.columns([5, 2, 1])
                col_n.write(f"**{i+1}.** {item['物品/费用名称']}")
                col_p.write(f"{item['金额（元）']:.2f} 元")
                col_d.button("🗑️", key=f"del_{i}_main", on_click=delete_excel_item, args=(i,), help="删除此项")
            total_preview = sum(x["金额（元）"] for x in st.session_state.excel_items)
            st.markdown(f"<div style='text-align:right;color:var(--primary);font-size:1.05rem;margin-top:10px;'>合计：{total_preview:.2f} 元</div>", unsafe_allow_html=True)
            c_dl, c_clr = st.columns([3, 1])
            with c_dl:
                excel_data, _ = build_excel(st.session_state.excel_items)
                st.download_button(
                    label="📥 下载 Excel 报销清单",
                    data=excel_data,
                    file_name=f"报销清单_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            with c_clr:
                st.button("清空全部", key="xl_clear_main", on_click=clear_excel_items, use_container_width=True)
        else:
            st.info("📂 暂无费用条目，请通过上方方式添加。")

    list_fragment()

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# 页面：流程引导
# ============================================================
elif page == "📋 流程引导":
    st.markdown('<div class="card"><div class="card-title">📋 报销流程引导</div>', unsafe_allow_html=True)
    st.caption("填写以下所有信息，点击「重新生成材料清单」即可获得完整报销指引。所有字段随时可修改。")

    demo_col1, demo_col2 = st.columns([1, 5])
    with demo_col1:
        if st.button("🚀 加载演示案例", key="demo_guide_top", use_container_width=True):
            load_demo_case()
            st.session_state.demo_loaded = True
    with demo_col2:
        if st.session_state.get("demo_loaded"):
            st.info("✅ 演示案例已加载，信息已自动填入下方表单。")

    @st.fragment
    def guide_fragment():
        guide_type = st.selectbox(
            "报销类型", list(GUIDE_TYPES.keys()),
            index=list(GUIDE_TYPES.keys()).index(st.session_state.guide_data["type"]),
            key="guide_type_frag"
        )
        st.session_state.guide_data["type"] = guide_type
        amount = st.session_state.guide_data["amount"]

        default_approver, default_submit_loc = get_auto_approver_submit(guide_type, amount)

        col1, col2 = st.columns(2)
        with col1:
            purpose = st.text_input("报销事由/活动名称", value=st.session_state.guide_data["purpose"], key="guide_purpose_frag")
            st.session_state.guide_data["purpose"] = purpose
            applicant = st.selectbox(
                "申请人身份", ["在校本科生", "在校研究生", "教职工", "博士生"],
                index=["在校本科生", "在校研究生", "教职工", "博士生"].index(st.session_state.guide_data["applicant"]),
                key="guide_applicant_frag"
            )
            st.session_state.guide_data["applicant"] = applicant
            amount_input = st.number_input("报销总金额（元）", min_value=0.0, step=10.0,
                                          value=st.session_state.guide_data["amount"], key="guide_amount_frag")
            st.session_state.guide_data["amount"] = amount_input
        with col2:
            dept = st.text_input("所在学院/部门", value=st.session_state.guide_data["dept"], key="guide_dept_frag")
            st.session_state.guide_data["dept"] = dept
            has_invoice = st.radio(
                "票据情况", ["✅ 有正规发票", "⚠️ 只有收据", "❌ 无任何票据"],
                index=["✅ 有正规发票", "⚠️ 只有收据", "❌ 无任何票据"].index(st.session_state.guide_data["has_invoice"]),
                key="guide_has_invoice_frag"
            )
            st.session_state.guide_data["has_invoice"] = has_invoice
            invoice_count = st.number_input("发票/收据张数", min_value=0, step=1,
                                           value=st.session_state.guide_data["invoice_count"], key="guide_invoice_count_frag")
            st.session_state.guide_data["invoice_count"] = invoice_count
        items_desc = st.text_area("费用明细说明（选填）", value=st.session_state.guide_data["items_desc"],
                                  height=80, key="guide_items_desc_frag")
        st.session_state.guide_data["items_desc"] = items_desc
        special_note = st.text_area("特殊说明（选填）", value=st.session_state.guide_data["special_note"],
                                    height=60, key="guide_special_note_frag")
        st.session_state.guide_data["special_note"] = special_note

        approver = st.text_input("审批人（姓名或职务）",
                                value=st.session_state.guide_data["approver"] if st.session_state.guide_data["approver"] else default_approver,
                                key="guide_approver_frag")
        st.session_state.guide_data["approver"] = approver
        submit_loc = st.text_input("提交地点/窗口",
                                  value=st.session_state.guide_data["submit_loc"] if st.session_state.guide_data["submit_loc"] else default_submit_loc,
                                  key="guide_submit_loc_frag")
        st.session_state.guide_data["submit_loc"] = submit_loc
        is_urgent = st.checkbox("⚡ 紧急报销（需加急处理）", value=st.session_state.guide_data["is_urgent"],
                               key="guide_is_urgent_frag")
        st.session_state.guide_data["is_urgent"] = is_urgent

        if st.button("📋 重新生成材料清单", key="guide_regenerate_frag", use_container_width=True):
            st.session_state.guide_generated = True

        if st.session_state.guide_generated:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.success("🎉 材料清单已生成！")

            output_html = generate_guide_output(st.session_state.guide_data)
            st.markdown(output_html, unsafe_allow_html=True)
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

            st.markdown("#### 📋 材料准备情况检查")
            st.caption("请勾选你已经准备好的材料，然后点击「检查材料完整性」")
            required_docs = GUIDE_TYPES[st.session_state.guide_data["type"]]["required_docs"]

            with st.form("material_check_form"):
                doc_checkboxes = {}
                for doc in required_docs:
                    default_val = doc in st.session_state.get("prepared_docs", [])
                    doc_checkboxes[doc] = st.checkbox(f"📄 {doc}", value=default_val, key=f"mat_{doc}")
                check_submitted = st.form_submit_button("检查材料完整性", use_container_width=True)

            if check_submitted:
                prepared_docs = [doc for doc, checked in doc_checkboxes.items() if checked]
                st.session_state.prepared_docs = prepared_docs
                _, missing_docs = check_missing_docs(required_docs, prepared_docs)
                st.session_state.missing_docs_count = len(missing_docs)

                progress_val = len(prepared_docs) / len(required_docs) if required_docs else 0
                pct_val = int(progress_val * 100)
                st.markdown(f'<div class="anim-progress-wrap"><div class="anim-progress-bar" style="width:{pct_val}%;animation:fillBar 0.8s ease-out forwards;"></div></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="mat-progress-text">材料准备进度：{len(prepared_docs)} / {len(required_docs)}</div>',
                           unsafe_allow_html=True)

                if prepared_docs:
                    st.success("你已准备：\n" + "\n".join([f"✅ {d}" for d in prepared_docs]))
                if missing_docs:
                    st.warning("仍需补充：\n" + "\n".join([f"❌ {d}" for d in missing_docs]))
                    st.error(f"🚨 风险提示：当前材料尚未完整，缺少 {len(missing_docs)} 项材料，可能导致报销申请被退回。建议提交前先补齐材料。")

                st.markdown("---")
                risk_result = calculate_reimbursement_risk(st.session_state.guide_data, len(missing_docs))
                st.session_state.risk_result = risk_result
                show_risk_card(risk_result)
            else:
                if st.session_state.risk_result is not None:
                    prepared_docs = st.session_state.get("prepared_docs", [])
                    missing_count = st.session_state.get("missing_docs_count", 0)
                    progress_val = len(prepared_docs) / len(required_docs) if required_docs else 0
                    pct_val = int(progress_val * 100)
                    st.markdown(f'<div class="anim-progress-wrap"><div class="anim-progress-bar" style="width:{pct_val}%;animation:fillBar 0.8s ease-out forwards;"></div></div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="mat-progress-text">材料准备进度：{len(prepared_docs)} / {len(required_docs)}</div>',
                               unsafe_allow_html=True)
                    show_risk_card(st.session_state.risk_result)

            st.markdown("---")
            st.markdown("#### 📝 报销申请说明")
            app_text = generate_application_text(st.session_state.guide_data)
            st.session_state.guide_application_text = app_text
            st.code(app_text, language="text")
            st.button("📋 复制文本", key="copy_app_text_btn",
                     on_click=lambda: st.write(""),
                     help="选中上方文本框中的文字后按 Ctrl+C 复制", use_container_width=True)

            st.markdown("---")
            st.markdown("#### 📥 下载报销材料")
            if st.session_state.risk_result is not None:
                prepared_docs = st.session_state.get("prepared_docs", [])
                _, missing_docs = check_missing_docs(required_docs, prepared_docs)
            else:
                prepared_docs = []
                missing_docs = required_docs

            if DOCX_AVAILABLE:
                try:
                    word_data = build_reimbursement_docx(
                        st.session_state.guide_data,
                        app_text,
                        st.session_state.risk_result or {"level": "低风险", "reasons": ["待检查"], "suggestions": ["请先检查材料完整性"]},
                        prepared_docs,
                        missing_docs,
                    )
                    st.download_button(
                        label="📄 下载 Word 版报销说明",
                        data=word_data,
                        file_name=f"报销申请说明_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.warning(f"Word 生成失败：{e}，请确认 python-docx 已安装。")
            else:
                st.info("💡 Word 导出功能需要安装 python-docx：`pip install python-docx`")

        else:
            st.info("👆 填写完信息后，点击「重新生成材料清单」按钮即可获得详细指引。")

    guide_fragment()
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# 页面：预算预警（专业面板布局）
# ============================================================
elif page == "⚠️ 预算预警":
    # 顶部标题栏
    col_title, col_meta = st.columns([2, 1])
    with col_title:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:10px;margin-bottom:14px;">
            <div style="font-size:1.3rem;font-weight:700;color:#1F4E8C;">⚠️ 预算预警与智能分析</div>
        </div>
        """, unsafe_allow_html=True)
    with col_meta:
        st.markdown(f"""
        <div style="text-align:right;font-size:0.72rem;color:#7A8FA0;padding-top:8px;">
            数据更新时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}
        </div>
        """, unsafe_allow_html=True)

    # ===== 左：输入表单 | 右：指标+图表 =====
    col_left, col_right = st.columns([1, 1.7])

    with col_left:
        # 表单区域（放卡片内）
        st.markdown('<div class="card" style="margin-bottom:16px;">', unsafe_allow_html=True)
        st.markdown("#### 📋 预算参数设置")

        with st.form("budget_form"):
            proj_name = st.text_input("预算项目名称", value=st.session_state.get("budget_name", ""), placeholder="例：五四团日活动")
            total_budget = st.number_input("预算总额（元）", min_value=0.0, step=100.0,
                                          value=st.session_state.get("budget_total", 10000.0))
            used_budget = st.number_input("已使用金额（元）", min_value=0.0, step=10.0,
                                         value=st.session_state.get("budget_used", 3500.0))
            deadline = st.date_input("预算截止日期", value=st.session_state.get("budget_deadline", date(date.today().year, 12, 31)))

            submitted = st.form_submit_button("📊 生成预算报告", use_container_width=True)

        if submitted:
            # 保存表单值到 session_state
            st.session_state.budget_name = proj_name
            st.session_state.budget_total = total_budget
            st.session_state.budget_used = used_budget
            st.session_state.budget_deadline = deadline

            # 计算
            total = float(total_budget)
            used = float(used_budget)
            remaining = total - used
            pct_used = (used / total * 100) if total > 0 else 0
            days_left = (deadline - date.today()).days
            days_total = (deadline - date(date.today().year, 1, 1)).days
            days_passed = max(days_total - days_left, 0)

            # 缓存结果
            st.session_state.budget_result = {
                "project_name": proj_name or "未命名项目",
                "total": total,
                "used": used,
                "remaining": remaining,
                "pct_used": pct_used,
                "days_left": days_left,
                "days_total": days_total,
                "days_passed": days_passed,
            }

        st.markdown('<div style="font-size:0.7rem;color:#A0B4C4;margin-top:10px;">数据仅用于当前计算，不做持久化存储</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col_right:
        result = st.session_state.get("budget_result")

        if result is None:
            # 无结果时显示引导
            st.markdown("""
            <div class="card" style="display:flex;align-items:center;justify-content:center;min-height:300px;">
                <div style="text-align:center;color:#A0B4C4;">
                    <div style="font-size:2.4rem;margin-bottom:10px;">📊</div>
                    <div style="font-size:0.95rem;">请在左侧填写预算参数后</div>
                    <div style="font-size:0.95rem;">点击「生成预算报告」查看分析</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            # 第一行：4个指标卡
            r = result
            days_left_val = max(r["days_left"], 0)
            met1, met2, met3, met4 = st.columns(4)
            met1.metric("预算总额", f"{r['total']:.0f} 元")
            met2.metric("已使用", f"{r['used']:.0f} 元", delta=f"{r['pct_used']:.1f}%")
            delta_color = "off" if r["remaining"] > 0 else "inverse"
            met3.metric("剩余预算", f"{r['remaining']:.2f} 元",
                        delta="充足" if r["remaining"] > r["total"]*0.3 else ("偏低" if r["remaining"] > 0 else "超支"),
                        delta_color=delta_color)
            met4.metric("距截止日期", f"{days_left_val} 天",
                        delta="时间充裕" if days_left_val > 30 else ("即将到期" if days_left_val > 0 else "已过期"),
                        delta_color="normal" if days_left_val > 14 else "inverse")

            # 第二行：图表 + 风险卡
            col_chart, col_risk = st.columns([1.2, 1])

            with col_chart:
                # 仪表盘
                color = "#27AE60" if r["pct_used"] < 60 else ("#F39C12" if r["pct_used"] < 85 else "#E74C3C")
                fig = go.Figure(go.Indicator(
                    mode="gauge+number+delta",
                    value=r["pct_used"],
                    number={"suffix": "%", "font": {"color": "#1F4E8C", "size": 30}},
                    delta={"reference": 60, "valueformat": ".1f", "increasing": {"color": "#E74C3C"}, "decreasing": {"color": "#27AE60"}},
                    gauge={
                        "axis": {"range": [0, 100], "tickcolor": "#7A8FA0", "tickfont": {"color": "#7A8FA0", "size": 10}},
                        "bar": {"color": color},
                        "bgcolor": "#F6FAFC",
                        "bordercolor": "#E5EEF3",
                        "steps": [
                            {"range": [0, 60], "color": "rgba(39,174,96,0.08)"},
                            {"range": [60, 85], "color": "rgba(243,156,18,0.08)"},
                            {"range": [85, 100], "color": "rgba(231,76,60,0.08)"},
                        ],
                        "threshold": {"line": {"color": "#1F4E8C", "width": 2}, "value": 80},
                    },
                    title={"text": f"{r['project_name']}<br>预算使用率", "font": {"color": "#7A8FA0", "size": 12}}
                ))
                fig.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    font={"family": "Noto Sans SC"},
                    height=230, margin=dict(t=30, b=0, l=20, r=20)
                )
                st.plotly_chart(fig, use_container_width=True)

                # 时间 vs 预算对比图
                time_pct = min((r["days_passed"] / max(r["days_total"], 1)) * 100, 100)
                st.markdown("**📅 时间 vs 预算进度**")
                fig2 = go.Figure()
                fig2.add_trace(go.Bar(x=[time_pct], y=["时间"], orientation='h', marker_color="#3498DB",
                                      name="已过时间", text=[f"{time_pct:.1f}%"], textposition='inside'))
                fig2.add_trace(go.Bar(x=[r["pct_used"]], y=["预算"], orientation='h', marker_color=color,
                                      name="已用预算", text=[f"{r['pct_used']:.1f}%"], textposition='inside'))
                fig2.update_layout(
                    barmode='overlay',
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    font={"color": "#7A8FA0", "size": 10},
                    xaxis={"range": [0, 100], "ticksuffix": "%", "gridcolor": "#E5EEF3"},
                    height=120, margin=dict(t=5, b=5, l=60, r=20),
                    showlegend=False,
                )
                st.plotly_chart(fig2, use_container_width=True)

            with col_risk:
                @st.fragment
                def budget_risk_fragment():
                    # 风险等级卡片
                    if r["pct_used"] > 85 or r["remaining"] < 0:
                        level, bg, border, icon, label = "高风险", "#FFF4F2", "#E74C3C", "🚨", "高风险"
                    elif r["pct_used"] > 60 or r["days_left"] < 14:
                        level, bg, border, icon, label = "中风险", "#FFF8F0", "#F39C12", "⚠️", "中风险"
                    else:
                        level, bg, border, icon, label = "低风险", "#F0FFF4", "#27AE60", "✅", "低风险"

                    risk_color = {"高风险": "#E74C3C", "中风险": "#F39C12", "低风险": "#27AE60"}
                    st.markdown(f"""
                    <div class="risk-card risk-{level}" style="background:{bg};border-left:5px solid {border};border-radius:12px;padding:16px 18px;margin-bottom:12px;">
                        <div class="risk-card-title">{icon} 预算风险：{label}</div>
                        <ul style="font-size:0.82rem;color:#333;padding-left:18px;">
                            <li>使用率 {r['pct_used']:.1f}%</li>
                            <li>剩余 {r['remaining']:.0f} 元</li>
                            <li>剩余 {days_left_val} 天</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)

                    # AI 建议按钮
                    if st.button("🤖 AI 智能分析建议", key="budget_ai_btn", use_container_width=True):
                        with st.spinner("AI 正在分析预算状况..."):
                            burn_rate_calc = r["used"] / max(r["days_passed"], 1)
                            projected_total_calc = r["used"] + burn_rate_calc * max(days_left_val, 0)
                            analysis_prompt = f"""
        你是高校财务专家。请根据以下数据给出简短、实用的建议（不超过3句话）：
        项目：{r['project_name']}
        预算总额：{r['total']} 元
        已使用：{r['used']} 元
        使用率：{r['pct_used']:.1f}%
        剩余天数：{days_left_val} 天
        日均消耗：{burn_rate_calc:.1f} 元/天
        到期预测总花费：{projected_total_calc:.0f} 元
        请回答：
        1. 是否存在超支风险？
        2. 给出 1-2 条具体调整建议（如控制支出、提前报销、申请预算调整等）。
        要求：简洁、直接、有用。
        """
                            try:
                                response = client.chat.completions.create(
                                    model="deepseek-chat",
                                    messages=[{"role": "user", "content": analysis_prompt}],
                                    temperature=0.5,
                                    max_tokens=250,
                                )
                                ai_response = response.choices[0].message.content
                                st.success(f"💡 **AI 建议**：{ai_response}")
                            except Exception as e:
                                st.error(f"AI 分析失败：{e}")
                budget_risk_fragment()

            # 趋势预测
            if r["days_passed"] > 0 and r["used"] > 0:
                burn_rate = r["used"] / r["days_passed"]
                if days_left_val > 0:
                    projected_total = r["used"] + burn_rate * days_left_val
                    overshoot = "超支" if projected_total > r["total"] else "未超支"
                    st.info(f"📈 **趋势预测**：当前日均消耗 {burn_rate:.1f} 元。若保持此速度，到期预计总花费 **{projected_total:.0f} 元**，将{overshoot} {abs(projected_total - r['total']):.0f} 元。")

            # 合规提醒
            st.markdown("**⚖️ 财务合规提醒**")
            if r["pct_used"] > 80 and days_left_val < 30:
                st.warning("⚠️ 预算使用率超过80%且剩余时间不足30天，请尽快处理已发生费用的报销，避免逾期失效。")
            elif r["pct_used"] > 80:
                st.warning("⚠️ 预算使用率超过80%，请严格控制后续支出。如确有需要，请提前申请预算调整。")
            elif days_left_val <= 7 and r["remaining"] > 0:
                st.warning(f"⏰ 距截止日期仅剩 {days_left_val} 天，请抓紧时间完成报销手续。")
            elif r["remaining"] < 0:
                st.error("🚨 预算已超支！请立即联系财务处申请追加预算或调整支出计划。")
            else:
                st.success("✅ 当前预算执行平稳，无重大合规风险。请继续保持合规报销。")

    # ===== 第三行：快捷入口 =====
    if result is not None:
        st.markdown('<div class="budget-section-title">🔗 快捷功能入口</div>', unsafe_allow_html=True)
        ec1, ec2, ec3 = st.columns(3)
        with ec1:
            if st.button("📋 材料缺失检查", key="goto_guide_mat", use_container_width=True):
                st.session_state.current_page = "📋 流程引导"
                st.session_state.nav_radio = "📋 流程引导"
        with ec2:
            if st.button("🚨 报销风险等级", key="goto_guide_risk", use_container_width=True):
                st.session_state.current_page = "📋 流程引导"
                st.session_state.nav_radio = "📋 流程引导"
        with ec3:
            if st.button("📝 报销申请说明", key="goto_guide_text", use_container_width=True):
                st.session_state.current_page = "📋 流程引导"
                st.session_state.nav_radio = "📋 流程引导"

# ═══════════════════════════════════════════════════════════════════════
# 页脚
# ═══════════════════════════════════════════════════════════════════════
st.markdown("""
<div style="text-align:center;color:#A0B4C4;font-size:0.7rem;margin-top:40px;padding:20px 0 10px;border-top:1px solid #E5EEF3;">
    校园财务问答助手 · 江西财经大学现代经济管理学院 · 比赛演示版
</div>
""", unsafe_allow_html=True)
